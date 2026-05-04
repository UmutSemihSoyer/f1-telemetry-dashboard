import docx

def read_docx(file_path):
    try:
        doc = docx.Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        
        # Also extract tables which usually contain the struct layouts
        for table in doc.tables:
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                full_text.append(" | ".join(row_data))
                
        print("\n".join(full_text))
    except Exception as e:
        print(f"Error reading docx: {e}")

if __name__ == "__main__":
    read_docx("Data Output from F1 22 v16.docx")
